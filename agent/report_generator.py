"""
Report Generator — Builds a clean .docx financial report using python-docx.

Sections:
  1. Title + Date header
  2. Market Overview table (Sensex, Nifty50)
  3. MCX Gold Price table (24K & 22K)
  4. Top Business News (1 best article)
  5. Disclaimer footer
"""

import os
import logging
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ── Colour palette ──────────────────────────────────────────────────────────
C_BLUE      = "1A56DB"   # header blue
C_GOLD      = "D97706"   # gold section
C_TEAL      = "0F766E"   # news section
C_GREEN     = "15803D"   # positive change
C_RED       = "DC2626"   # negative change
C_GRAY      = "6B7280"   # muted text
C_WHITE     = "FFFFFF"
C_LIGHT_BG  = "F3F4F6"   # alternating row background


# ── Helpers ─────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Apply a solid background colour to a table cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold=False, color_hex: str = "000000",
                   size_pt: int = 10, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    """Set cell text with formatting."""
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )


def _add_section_heading(doc: Document, text: str, color_hex: str) -> None:
    """Add a styled level-1 heading."""
    heading = doc.add_heading(text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16),
        )
        run.font.size = Pt(14)


def _add_divider(doc: Document) -> None:
    """Add a thin horizontal rule paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("─" * 80)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)


# ── Main generator ───────────────────────────────────────────────────────────

def generate_report(
    sensex_data: dict,
    nifty_data: dict,
    gold_data: dict,
    news_data: list,
    ai_commentary: str = "",  # kept for signature compatibility, not used
) -> str:
    """
    Build a .docx financial report and save it to /tmp.

    Returns:
        Absolute file path of the saved .docx file.
    """
    doc = Document()

    # ── Page margins ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Title ────────────────────────────────────────────────────────────────
    title_para = doc.add_heading("📊  Daily Financial Market Report", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.bold = True

    # Subtitle / date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(
        f"Report generated on  {datetime.now().strftime('%A, %d %B %Y  |  %I:%M %p IST')}"
    )
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    date_run.italic = True

    _add_divider(doc)

    # ── Market Overview ──────────────────────────────────────────────────────
    _add_section_heading(doc, "📈  Market Overview", C_BLUE)

    market_table = doc.add_table(rows=1, cols=5)
    market_table.style = "Table Grid"
    market_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_headers = ["Index", "Current Price (₹)", "Prev. Close (₹)", "Change (₹)", "% Change"]
    col_widths   = [Inches(1.6), Inches(1.5), Inches(1.5), Inches(1.2), Inches(1.1)]

    # Set column widths
    for i, cell in enumerate(market_table.rows[0].cells):
        cell.width = col_widths[i]

    # Header row
    hdr_row = market_table.rows[0]
    for i, hdr in enumerate(col_headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, C_BLUE)
        _set_cell_text(cell, hdr, bold=True, color_hex=C_WHITE, size_pt=10)

    # Data rows
    for idx, data in enumerate([sensex_data, nifty_data]):
        if "error" in data:
            row = market_table.add_row()
            _set_cell_text(row.cells[0], data.get("name", "N/A"))
            _set_cell_text(row.cells[1], "Data unavailable", color_hex=C_GRAY)
            continue

        row = market_table.add_row()
        bg = C_LIGHT_BG if idx % 2 else C_WHITE

        change   = data.get("change", 0)
        pct      = data.get("percent_change", 0)
        is_up    = change >= 0
        val_color = C_GREEN if is_up else C_RED

        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)
        _set_cell_bg(row.cells[2], bg)
        _set_cell_bg(row.cells[3], "E8F5E9" if is_up else "FFEBEE")
        _set_cell_bg(row.cells[4], "E8F5E9" if is_up else "FFEBEE")

        _set_cell_text(row.cells[0], data.get("name", ""), bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_text(row.cells[1], f"₹{data.get('current_price', 0):,.2f}")
        _set_cell_text(row.cells[2], f"₹{data.get('previous_close', 0):,.2f}")
        _set_cell_text(row.cells[3],
                       f"{'+' if is_up else ''}{change:,.2f}",
                       bold=True, color_hex=val_color)
        _set_cell_text(row.cells[4],
                       f"{data.get('direction', '')} {'+' if is_up else ''}{pct:.2f}%",
                       bold=True, color_hex=val_color)

    doc.add_paragraph()

    # ── Gold Price ───────────────────────────────────────────────────────────
    _add_section_heading(doc, "🥇  MCX Gold Price  (National Rate)", C_GOLD)

    if "error" not in gold_data:
        gold_table = doc.add_table(rows=1, cols=5)
        gold_table.style = "Table Grid"
        gold_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        g_headers = ["Purity", "Per Gram (₹)", "Per 10g (₹)", "Change / 10g (₹)", "% Change"]
        g_widths   = [Inches(1.2), Inches(1.4), Inches(1.4), Inches(1.6), Inches(1.2)]

        for i, cell in enumerate(gold_table.rows[0].cells):
            cell.width = g_widths[i]

        g_hdr_row = gold_table.rows[0]
        for i, hdr in enumerate(g_headers):
            cell = g_hdr_row.cells[i]
            _set_cell_bg(cell, C_GOLD)
            _set_cell_text(cell, hdr, bold=True, color_hex=C_WHITE, size_pt=10)

        gold_change = gold_data.get("change_per_10g", 0)
        gold_pct    = gold_data.get("percent_change", 0)
        gold_up     = gold_change >= 0
        gold_vc     = C_GREEN if gold_up else C_RED

        for purity, pg, p10 in [
            ("24K (Pure Gold)",
             gold_data.get("price_per_gram_inr", "N/A"),
             gold_data.get("price_per_10g_inr", "N/A")),
            ("22K (Jewellery)",
             gold_data.get("price_per_gram_inr_22k", "N/A"),
             round(gold_data.get("price_per_gram_inr_22k", 0) * 10, 2)),
        ]:
            row = gold_table.add_row()
            _set_cell_text(row.cells[0], purity, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
            _set_cell_text(row.cells[1], f"₹{pg:,.2f}" if isinstance(pg, float) else str(pg))
            _set_cell_text(row.cells[2], f"₹{p10:,.2f}" if isinstance(p10, float) else str(p10))
            _set_cell_bg(row.cells[3], "E8F5E9" if gold_up else "FFEBEE")
            _set_cell_bg(row.cells[4], "E8F5E9" if gold_up else "FFEBEE")
            _set_cell_text(row.cells[3],
                           f"{'+' if gold_up else ''}{gold_change:,.2f}",
                           bold=True, color_hex=gold_vc)
            _set_cell_text(row.cells[4],
                           f"{gold_data.get('direction', '')} {'+' if gold_up else ''}{gold_pct:.2f}%",
                           bold=True, color_hex=gold_vc)

        # Reference note
        ref_para = doc.add_paragraph()
        ref_run = ref_para.add_run(
            f"📌  Reference: COMEX Gold = ${gold_data.get('usd_price_per_oz', 'N/A')} / oz  |  "
            f"USD/INR = ₹{gold_data.get('usd_inr_rate', 'N/A')}  |  "
            f"As of: {gold_data.get('as_of', '')}"
        )
        ref_run.font.size = Pt(8.5)
        ref_run.font.italic = True
        ref_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    else:
        err_para = doc.add_paragraph(f"⚠️  Gold data unavailable: {gold_data.get('error')}")
        err_para.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    doc.add_paragraph()

    # ── Top Business News (1 best article) ───────────────────────────────────
    _add_section_heading(doc, "📰  Top Business News  (India — Today)", C_TEAL)

    valid_news = [a for a in news_data if isinstance(a, dict) and "error" not in a]

    if not valid_news:
        doc.add_paragraph("⚠️  No business news available at this time.")
    else:
        # Pick the single best article (first valid one from the ranked list)
        article = valid_news[0]

        # Headline
        headline_para = doc.add_paragraph()
        headline_run = headline_para.add_run(article.get("title", ""))
        headline_run.bold = True
        headline_run.font.size = Pt(11)
        headline_run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        headline_para.paragraph_format.space_before = Pt(4)

        # Description
        desc = article.get("description", "").strip()
        if desc:
            desc_para = doc.add_paragraph(desc)
            desc_para.runs[0].font.size = Pt(10)
            desc_para.runs[0].font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            desc_para.paragraph_format.space_before = Pt(2)
            desc_para.paragraph_format.space_after  = Pt(2)

        # Source & date
        meta_para = doc.add_paragraph(
            f"Source: {article.get('source', 'N/A')}   |   {article.get('published_at', '')}"
        )
        meta_para.runs[0].font.size = Pt(9)
        meta_para.runs[0].font.italic = True
        meta_para.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
        meta_para.paragraph_format.space_after = Pt(6)

        # Read more link
        url = article.get("url", "")
        if url:
            link_para = doc.add_paragraph(f"Read more: {url}")
            link_para.runs[0].font.size = Pt(9)
            link_para.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    # ── Footer ──────────────────────────────────────────────────────────────

    # ── Save ─────────────────────────────────────────────────────────────────
    filename  = f"Aditya_Rajput_{datetime.now().strftime('%Y-%m-%d')}.docx"
    save_dir  = os.getenv("REPORT_SAVE_DIR", "/tmp")
    filepath  = os.path.join(save_dir, filename)
    doc.save(filepath)
    logger.info("Report saved to: %s", filepath)
    return filepath
